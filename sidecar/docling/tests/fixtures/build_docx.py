"""Deterministic table_invoice.docx generator (T-04-18's Python half).

`docx` (python-docx) is already installed -- msword_backend.py depends on it -- so no new
requirements.txt line is needed. Document.save() writes each OOXML package part through
zipfile.writestr(name, blob), which stamps the CURRENT wall-clock time into every entry's
date_time field: two saves seconds apart never byte-match. Re-zipping every member with a
fixed date_time removes that.

Compression is left on (ZIP_DEFLATED, compresslevel=9) rather than stored: python-docx's
built-in template carries ~800 KB of unused default style XML that ZIP_STORED would commit
verbatim, and this generator only ever runs inside the pinned docling:test image (one zlib
build), so DEFLATE's output for fixed input is stable across every run that matters here --
a cross-machine risk this project's single-image CI never actually exercises. Sorted member
order makes the compare independent of dict iteration order too.

Mirrors fixtures_test.go's -update idiom: `python build_docx.py --update` regenerates the
committed file; a bare run byte-compares and exits non-zero on drift.
"""

import io
import sys
import zipfile
from pathlib import Path

from docx import Document

FIXTURE = Path(__file__).parent / "table_invoice.docx"

# Same header/body text as internal/extraction/fixtures_test.go's fxBuildTable, for
# cross-format parity between the two fixtures this subtask commissions.
HEADER = ["Description", "Qty", "Unit Price", "Total"]
BODY_ROWS = [
    ["Widget", "2", "500.00", "1000.00"],
    ["Gadget", "1", "500.00", "500.00"],
]


def build() -> bytes:
    doc = Document()
    doc.add_paragraph("INVOICE")
    table = doc.add_table(rows=1 + len(BODY_ROWS), cols=len(HEADER))
    for col, text in enumerate(HEADER):
        table.cell(0, col).text = text
    for row_idx, row in enumerate(BODY_ROWS, start=1):
        for col, text in enumerate(row):
            table.cell(row_idx, col).text = text

    raw = io.BytesIO()
    doc.save(raw)
    return _normalise_zip(raw.getvalue())


def _normalise_zip(data: bytes) -> bytes:
    src = zipfile.ZipFile(io.BytesIO(data))
    out = io.BytesIO()
    with zipfile.ZipFile(out, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=9) as dst:
        for name in sorted(src.namelist()):
            info = zipfile.ZipInfo(name, date_time=(1980, 1, 1, 0, 0, 0))
            dst.writestr(info, src.read(name), compress_type=zipfile.ZIP_DEFLATED)
    return out.getvalue()


if __name__ == "__main__":
    built = build()
    if "--update" in sys.argv:
        FIXTURE.write_bytes(built)
        print(f"wrote {FIXTURE} ({len(built)} bytes)")
    else:
        existing = FIXTURE.read_bytes()
        if existing != built:
            sys.exit(
                f"{FIXTURE} does not match its generator: {len(existing)} byte(s) on disk, "
                f"{len(built)} regenerated -- run `python build_docx.py --update`"
            )
        print("table_invoice.docx matches its generator")
